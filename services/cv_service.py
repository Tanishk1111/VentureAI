import os
import PyPDF2
import docx
import io
import fitz  # PyMuPDF for better PDF extraction
from utils.gemini_utils import generate_content, generate_cv_questions_enhanced

def extract_text_from_cv(cv_path):
    """Extract text content from a CV file (PDF or DOCX) with improved accuracy"""
    text = ""
    file_extension = os.path.splitext(cv_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            # Process PDF file using PyMuPDF for better extraction
            try:
                with fitz.open(cv_path) as pdf:
                    for page in pdf:
                        text += page.get_text()
            except ImportError:
                # Fallback to PyPDF2 if PyMuPDF is not available
                with open(cv_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        text += page.extract_text()
        
        elif file_extension in ['.docx', '.doc']:
            # Process Word document
            doc = docx.Document(cv_path)
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        
        else:
            # For text files or other formats
            with open(cv_path, 'r', encoding='utf-8', errors='ignore') as file:
                text = file.read()
        
        # Clean up the extracted text
        text = clean_cv_text(text)
        return text
    except Exception as e:
        print(f"Error extracting text from CV: {e}")
        return ""

def clean_cv_text(text):
    """Clean and format extracted CV text"""
    # Replace multiple newlines with a single newline
    import re
    text = re.sub(r'\n+', '\n', text)
    
    # Replace multiple spaces with a single space
    text = re.sub(r' +', ' ', text)
    
    # Remove special characters that might interfere with analysis
    text = text.replace('\x0c', ' ').replace('\t', ' ')
    
    return text.strip()

def save_uploaded_cv(file_content, filename, upload_dir="uploads"):
    """Save an uploaded CV and return the file path"""
    os.makedirs(upload_dir, exist_ok=True)
    
    file_path = os.path.join(upload_dir, filename)
    with open(file_path, 'wb') as f:
        f.write(file_content)
    
    return file_path

def generate_cv_questions(cv_text, num_questions=5):
    """Generate interview questions based on CV content using enhanced Gemini prompting"""
    # Use the enhanced CV question generation with better prompting
    all_questions = generate_cv_questions_enhanced(cv_text)
    
    # Return up to the requested number of questions
    return all_questions[:num_questions]
